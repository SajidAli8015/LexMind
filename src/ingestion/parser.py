"""
Document Parser for LexMind
Converts raw legal documents into structured elements
preserving document hierarchy for legal-aware chunking.
"""

import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional
from loguru import logger


# ─── Data Classes ────────────────────────────────────────────

@dataclass
class DocumentElement:
    """
    Represents one piece of content from a legal document.

    Examples:
        DocumentElement(text="ARTICLE 47 TERMINATION",
                       element_type="heading", level=2)
        DocumentElement(text="Either party may terminate...",
                       element_type="paragraph", level=0)
        DocumentElement(text="47.1 Notice must be given...",
                       element_type="heading", level=3)
    """
    text: str
    element_type: str        # heading, paragraph, table, list_item
    level: int = 0           # 0=normal, 1=Part/Chapter,
                             # 2=Article/Section, 3=Sub-clause
    page_number: int = 0
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """
    Represents a fully parsed legal document with all its elements.

    Contains all DocumentElements from the file plus document-level
    metadata like page count and language.
    """
    file_path: str
    file_name: str
    file_type: str           # pdf, docx, txt
    elements: List[DocumentElement] = field(default_factory=list)
    total_pages: int = 0
    language: str = "en"
    metadata: dict = field(default_factory=dict)

    def element_count(self) -> int:
        """Total number of elements parsed from the document."""
        return len(self.elements)

    def get_headings(self) -> List[DocumentElement]:
        """Return only heading elements (Articles, Parts, Sub-clauses)."""
        return [e for e in self.elements if e.element_type == "heading"]

    def get_full_text(self) -> str:
        """Return all element text joined as one string."""
        return "\n".join(e.text for e in self.elements if e.text.strip())

    def get_articles(self) -> List[DocumentElement]:
        """Return only Article-level headings (level 2)."""
        return [e for e in self.elements
                if e.element_type == "heading" and e.level == 2]


# ─── Parser Class ────────────────────────────────────────────

class DocumentParser:
    """
    Parses legal documents into structured DocumentElements.

    Tries Docling first for best quality structure extraction.
    Falls back to simpler parsers if Docling is unavailable.

    Supported formats: PDF, DOCX, TXT
    """

    def __init__(self):
        """Initialize parser and check Docling availability."""
        self.docling_available = False
        self._converter = None

        try:
            from docling.document_converter import DocumentConverter
            self._converter = DocumentConverter()
            self.docling_available = True
            logger.info("DocumentParser initialized with Docling")
        except ImportError:
            logger.warning(
                "Docling not available — using fallback parsers. "
                "Install with: pip install docling"
            )
        except Exception as e:
            logger.warning(f"Docling initialization failed: {e} — using fallbacks")

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a legal document file into structured elements.

        Args:
            file_path: Path to the document (PDF, DOCX, or TXT)

        Returns:
            ParsedDocument with all elements extracted

        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If file type is not supported
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = path.suffix.lower().lstrip(".")
        logger.info(f"Parsing {path.name} (type: {file_type})")

        if file_type == "txt":
            doc = self._parse_txt(file_path)
        elif file_type in ("pdf", "docx", "pptx", "html"):
            if self.docling_available:
                try:
                    doc = self._parse_with_docling(file_path, file_type)
                except Exception as e:
                    logger.warning(
                        f"Docling failed: {e} — trying fallback parser"
                    )
                    doc = self._parse_fallback(file_path, file_type)
            else:
                doc = self._parse_fallback(file_path, file_type)
        else:
            raise ValueError(
                f"Unsupported file type: .{file_type}. "
                f"Supported: pdf, docx, txt, pptx, html"
            )

        logger.info(
            f"Parsed {path.name}: "
            f"{doc.element_count()} elements, "
            f"{len(doc.get_headings())} headings, "
            f"{len(doc.get_articles())} articles"
        )
        return doc

    def _parse_with_docling(
        self, file_path: str, file_type: str
    ) -> ParsedDocument:
        """
        Parse document using Docling for best structure quality.
        Docling preserves headings, tables, and document hierarchy.
        """
        from docling.document_converter import DocumentConverter

        path = Path(file_path)
        result = self._converter.convert(file_path)
        doc_obj = result.document

        elements = []
        page_count = 0

        # Export to markdown to get structured text
        markdown_text = doc_obj.export_to_markdown()

        # Parse the markdown structure
        for line in markdown_text.split("\n"):
            line = line.strip()
            if not line:
                continue

            # Markdown headings: # ## ###
            if line.startswith("### "):
                text = line[4:].strip()
                level = self._detect_heading_level(text) or 3
                elements.append(DocumentElement(
                    text=text,
                    element_type="heading",
                    level=level
                ))
            elif line.startswith("## "):
                text = line[3:].strip()
                level = self._detect_heading_level(text) or 2
                elements.append(DocumentElement(
                    text=text,
                    element_type="heading",
                    level=level
                ))
            elif line.startswith("# "):
                text = line[2:].strip()
                elements.append(DocumentElement(
                    text=text,
                    element_type="heading",
                    level=1
                ))
            elif line.startswith("|"):
                # Markdown table row
                elements.append(DocumentElement(
                    text=line,
                    element_type="table",
                    level=0
                ))
            elif line.startswith("- ") or line.startswith("* "):
                elements.append(DocumentElement(
                    text=line[2:].strip(),
                    element_type="list_item",
                    level=0
                ))
            else:
                # Check if this looks like a legal heading
                level = self._detect_heading_level(line)
                if level > 0:
                    elements.append(DocumentElement(
                        text=line,
                        element_type="heading",
                        level=level
                    ))
                else:
                    elements.append(DocumentElement(
                        text=line,
                        element_type="paragraph",
                        level=0
                    ))

        # Try to get page count from docling metadata
        try:
            page_count = len(doc_obj.pages) if hasattr(doc_obj, 'pages') else 0
        except Exception:
            page_count = 0

        return ParsedDocument(
            file_path=str(file_path),
            file_name=path.name,
            file_type=file_type,
            elements=elements,
            total_pages=page_count,
            metadata={"parser": "docling"}
        )

    def _parse_txt(self, file_path: str) -> ParsedDocument:
        """
        Parse a plain text file.
        Uses pattern matching to detect headings from text patterns.
        """
        path = Path(file_path)

        # Try UTF-8 first, fall back to latin-1
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
            logger.warning(f"Used latin-1 fallback encoding for {path.name}")

        elements = []
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            level = self._detect_heading_level(line)
            if level > 0:
                elements.append(DocumentElement(
                    text=line,
                    element_type="heading",
                    level=level
                ))
            elif line.startswith("-") or line.startswith("•"):
                elements.append(DocumentElement(
                    text=line.lstrip("-•").strip(),
                    element_type="list_item",
                    level=0
                ))
            else:
                elements.append(DocumentElement(
                    text=line,
                    element_type="paragraph",
                    level=0
                ))

        return ParsedDocument(
            file_path=str(file_path),
            file_name=path.name,
            file_type="txt",
            elements=elements,
            total_pages=1,
            metadata={"parser": "txt_fallback"}
        )

    def _parse_fallback(self, file_path: str, file_type: str) -> ParsedDocument:
        """
        Fallback parser for PDF and DOCX when Docling unavailable.
        Less accurate than Docling but works without heavy dependencies.
        """
        if file_type == "pdf":
            return self._parse_pdf_fallback(file_path)
        elif file_type == "docx":
            return self._parse_docx_fallback(file_path)
        else:
            raise ValueError(
                f"No fallback parser for .{file_type}. "
                f"Please install docling: pip install docling"
            )

    def _parse_pdf_fallback(self, file_path: str) -> ParsedDocument:
        """
        Basic PDF parser using pypdf.
        Extracts text page by page with heading detection.
        """
        try:
            import pypdf
        except ImportError:
            raise ImportError(
                "pypdf not installed. Install with: pip install pypdf"
            )

        path = Path(file_path)
        elements = []

        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)

            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    level = self._detect_heading_level(line)
                    elem_type = "heading" if level > 0 else "paragraph"
                    elements.append(DocumentElement(
                        text=line,
                        element_type=elem_type,
                        level=level,
                        page_number=page_num
                    ))

        return ParsedDocument(
            file_path=str(file_path),
            file_name=path.name,
            file_type="pdf",
            elements=elements,
            total_pages=total_pages,
            metadata={"parser": "pypdf_fallback"}
        )

    def _parse_docx_fallback(self, file_path: str) -> ParsedDocument:
        """
        Basic DOCX parser using python-docx.
        Uses paragraph styles to detect headings.
        """
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        path = Path(file_path)
        document = docx.Document(file_path)
        elements = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if "Heading 1" in style_name:
                level = 1
                elem_type = "heading"
            elif "Heading 2" in style_name:
                level = 2
                elem_type = "heading"
            elif "Heading 3" in style_name:
                level = 3
                elem_type = "heading"
            else:
                # Fall back to pattern detection
                level = self._detect_heading_level(text)
                elem_type = "heading" if level > 0 else "paragraph"

            elements.append(DocumentElement(
                text=text,
                element_type=elem_type,
                level=level
            ))

        return ParsedDocument(
            file_path=str(file_path),
            file_name=path.name,
            file_type="docx",
            elements=elements,
            total_pages=0,
            metadata={"parser": "python_docx_fallback"}
        )

    def _detect_heading_level(self, text: str) -> int:
        """
        Detect if a line of text is a legal heading and what level.

        Returns:
            0 = not a heading (regular paragraph)
            1 = Part or Chapter level  (PART I, CHAPTER 3)
            2 = Article or Section     (ARTICLE 47, SECTION 3)
            3 = Sub-clause             (47.1, 3.2.1, (a), (i))

        Examples:
            "PART I GENERAL PROVISIONS"      → 1
            "ARTICLE 47 TERMINATION"         → 2
            "Article 3 Definitions"          → 2
            "47.1 Notice must be given"      → 3
            "3.2.1 Sub-clause content"       → 3
            "(a) The following applies"      → 3
            "Either party may terminate..."  → 0
        """
        if not text or len(text.strip()) < 2:
            return 0

        text_stripped = text.strip()
        text_upper = text_stripped.upper()

        # Level 1: PART or CHAPTER
        if re.match(
            r'^(PART|CHAPTER)\s+[IVX\d]', text_upper
        ):
            return 1

        # Level 2: ARTICLE or SECTION (case insensitive)
        if re.match(
            r'^(ARTICLE|SECTION|CLAUSE|SCHEDULE)\s+\d+',
            text_upper
        ):
            return 2

        # Level 2: All caps short line (likely a section title)
        if (text_stripped.isupper()
                and 3 < len(text_stripped) < 80
                and not text_stripped.endswith(".")):
            return 2

        # Level 3: Numbered sub-clause like 47.1 or 3.2.1
        if re.match(r'^\d+\.\d+(\.\d+)*\s', text_stripped):
            return 3

        # Level 3: Lettered sub-clause like (a) or (i)
        if re.match(r'^\([a-z]{1,3}\)\s', text_stripped):
            return 3

        # Level 3: Roman numeral sub-clause like (i) (ii) (iii)
        if re.match(r'^\([ivxlIVXL]+\)\s', text_stripped):
            return 3

        return 0


# ─── Convenience Function ────────────────────────────────────

def parse_document(file_path: str) -> ParsedDocument:
    """
    Convenience function to parse a document in one line.

    Usage:
        from ingestion.parser import parse_document
        doc = parse_document("contract.pdf")
        print(doc.element_count())
    """
    parser = DocumentParser()
    return parser.parse(file_path)


# ─── Quick Test ──────────────────────────────────────────────

if __name__ == "__main__":
    import tempfile, os

    sample = """PART I - GENERAL PROVISIONS

ARTICLE 1 DEFINITIONS
In this Agreement the following terms shall have meanings:

1.1 Agreement means this contract dated January 2024.
1.2 Party means either ABC Corporation or XYZ Limited.

ARTICLE 2 OBLIGATIONS
ABC Corporation shall perform the following:

2.1 Deliver services within 30 days of the Effective Date.
2.2 Maintain confidentiality of all shared information.

ARTICLE 3 PAYMENT TERMS
XYZ Limited shall pay PKR 5000000 within 30 days.

3.1 Late payments attract 2 percent monthly interest.
3.2 Disputes must be raised within 14 days of invoice.
"""
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as f:
        f.write(sample)
        tmp = f.name

    doc = parse_document(tmp)
    print(f"Elements : {doc.element_count()}")
    print(f"Headings : {len(doc.get_headings())}")
    print(f"Articles : {len(doc.get_articles())}")
    print()
    for i, e in enumerate(doc.elements, 1):
        label = f"[{e.element_type.upper():9}  L{e.level}]"
        preview = e.text[:55] + "..." if len(e.text) > 55 else e.text
        print(f"  {i:2}. {label}  {preview}")

    os.unlink(tmp)
    print("\nParser test PASSED" if doc.element_count() > 0 else "\nParser test FAILED")
